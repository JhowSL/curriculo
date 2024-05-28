from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

# Create a new Document
doc = Document()


# Define a function to add headings with custom formatting
def add_heading(text, level, font_size, bold=True, color=None):
    paragraph = doc.add_heading(text, level=level)
    run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


# Define a function to add paragraphs with custom formatting
def add_paragraph(text, font_size, color=None, alignment=None):
    paragraph = doc.add_paragraph(text)
    run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        paragraph.alignment = alignment


# Add Name
add_heading("Jonathan Souza Lima", 0, 24, bold=True)

# Add Contact Information
doc.add_heading("CONTACT", level=2)
add_paragraph(
    "São Paulo\n+55 11 9 8520-4061\njon.1lima@hotmail.com\nLinkedIn Profile: http://www.linkedin.com/in/jonathan-souza-lima-354967101",
    12,
)

# Add Professional Summary
doc.add_heading("PROFESSIONAL SUMMARY", level=2)
add_paragraph(
    "I am seeking a Junior Developer opportunity where I can apply and enhance my knowledge while actively contributing to process automation and front-end solution development. I have solid experience in web programming, focusing on technologies such as React, TypeScript, and Next.js. My dedication is clear through my participation in continuous improvements and the exploration of modern technologies. ",
    12,
)

# Add Work Experience
doc.add_heading("WORK EXPERIENCE", level=2)

# Experience 1
add_heading("Systems Development Intern / Vivo (Telefônica Brasil)", 3, 12, bold=True)
add_paragraph("April 2022 – April 2024", 10)
add_paragraph(
    "- Developed ShellScript automation solutions, improving operational efficiency.\n"
    "- Created detailed documentation and manuals to facilitate system usage.\n"
    "- Participated in the implementation of Backstage.io for infrastructure management.\n"
    "- Built front-end interfaces with React, TypeScript, and Next.js, ensuring intuitive user experiences.\n"
    "- Used ChatGPT to develop chatbots and virtual assistants, including the 'Hercules' project focusing on visual aspects and OpenAI integration.",
    10,
)

# Experience 2
add_heading("IT Intern / Store SP Outlet", 3, 12, bold=True)
add_paragraph("June 2021 – February 2022", 10)
add_paragraph(
    "- Maintained IT infrastructure, configured backup routines (EaseUS), planned and executed network cabling, assembled and maintained machines, and provided remote technical support via Anydesk.",
    10,
)

# Add Education
doc.add_heading("EDUCATION", level=2)
add_paragraph(
    "July 2021 - July 2024\nAnalysis and Systems Development / Faculdade Impacta de Tecnologia",
    12,
)

# Add Skills
doc.add_heading("SKILLS", level=2)
add_paragraph(
    "GitHub, Docker, front-end programming with TypeScript, Next.js, JWT, shadcn/ui, Tailwind CSS. Easily adapts to teamwork. Performance and curiosity in new activities. Initiative-taking",
    12,
)

# Save the document
file_path = "Jonathan_Souza_Lima_CV_American_Style.docx"
doc.save(file_path)

print(f"Document saved as {file_path}")
